"""Real text extraction for .pdf/.docx/.xlsx attachments (Nango live path).

Form-fields + text-layer only, pure Python — no OS-level OCR dependency. A
genuinely scanned/image-only PDF (no form fields, no text layer) honestly
extracts no text, same as an unsupported format; the caller falls back to
storing the raw bytes rather than fabricating content.
"""

from __future__ import annotations

import io
from pathlib import Path


def extract_text(filename: str, raw_bytes: bytes) -> str | None:
    """Best-effort real extraction. Never raises — a corrupt/unexpected file just
    yields ``None`` so the caller's existing utf-8-or-base64 fallback still runs."""
    suffix = Path(filename).suffix.lower()
    try:
        if suffix == ".pdf":
            text = _extract_pdf(raw_bytes)
        elif suffix == ".docx":
            text = _extract_docx(raw_bytes)
        elif suffix == ".xlsx":
            text = _extract_xlsx(raw_bytes)
        else:
            return None
    except Exception:
        return None
    return text if text and text.strip() else None


def _extract_pdf(raw_bytes: bytes) -> str | None:
    """Combines the page text layer with form-field values — a fillable form's
    static text (titles, headers, labels the classifier looks for, e.g. "ACORD")
    lives in the content stream, while the actual filled-in values live on the
    widget annotations and never appear in ``extract_text()``. Neither source
    alone is complete; real ACORD PDFs need both."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw_bytes))
    parts: list[str] = []

    text_layer = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    if text_layer:
        parts.append(text_layer)

    fields = reader.get_fields()
    if fields:
        lines = [
            f"{(f.alternate_name or f.name)}: {f.value}"
            for f in fields.values()
            if f.value not in (None, "")
        ]
        if lines:
            parts.append("\n".join(lines))

    return "\n".join(parts) if parts else None


def _extract_docx(raw_bytes: bytes) -> str | None:
    from docx import Document

    doc = Document(io.BytesIO(raw_bytes))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if not any(cells):
                continue
            lines.append(f"{cells[0]}: {cells[1]}" if len(cells) == 2 else " | ".join(cells))

    return "\n".join(lines)


def _extract_xlsx(raw_bytes: bytes) -> str | None:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(raw_bytes), data_only=True, read_only=True)
    lines: list[str] = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if not cells:
                continue
            lines.append(f"{cells[0]}: {cells[1]}" if len(cells) == 2 else " | ".join(cells))
    return "\n".join(lines)
