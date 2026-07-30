"""Proves core/ingestion/document_text.py against REAL generated PDF/DOCX/XLSX
files (built with reportlab/python-docx/openpyxl, the same libraries used to
read them) — not mocks of the extraction step. This is the real-document text
extraction added for the Nango live connector's Market Matching path.
"""

from __future__ import annotations

import io

import pytest
from docx import Document
from openpyxl import Workbook
from reportlab.pdfgen import canvas

from core.ingestion.document_text import extract_text


def _build_form_pdf() -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=(400, 300))
    c.drawString(20, 260, "ACORD 125 - COMMERCIAL INSURANCE APPLICATION")
    form = c.acroForm
    form.textfield(
        name="named_insured", tooltip="Named Insured",
        x=20, y=200, width=250, height=20,
        value="Delta Electric Services LLC", forceBorder=True,
    )
    form.textfield(
        name="class_code", tooltip="Class Code",
        x=20, y=160, width=250, height=20,
        value="contractors - electrical", forceBorder=True,
    )
    c.save()
    return buf.getvalue()


def _build_plain_text_pdf() -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=(400, 300))
    c.drawString(20, 260, "Loss Run Summary")
    c.drawString(20, 240, "Total Incurred 5yr: $20,700")
    c.save()
    return buf.getvalue()


def _build_blank_pdf() -> bytes:
    buf = io.BytesIO()
    canvas.Canvas(buf, pagesize=(400, 300)).save()
    return buf.getvalue()


def _build_docx() -> bytes:
    doc = Document()
    doc.add_paragraph("Financial Statement Summary")
    doc.add_paragraph("Annual Revenue: $5,800,000")
    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0]
    row.cells[0].text = "Total Assets"
    row.cells[1].text = "$2,400,000"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_xlsx() -> bytes:
    wb = Workbook()
    kv_sheet = wb.active
    kv_sheet.title = "Summary"
    kv_sheet.append(["Total Insurable Value", 8_500_000])
    wide_sheet = wb.create_sheet("Locations")
    wide_sheet.append(["Address", "Building Value", "Construction Type"])
    wide_sheet.append(["100 Main St", 1_200_000, "Frame"])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def test_pdf_form_fields_extracted_as_key_value() -> None:
    text = extract_text("acord_application.pdf", _build_form_pdf())
    assert text is not None
    assert "Named Insured: Delta Electric Services LLC" in text
    assert "Class Code: contractors - electrical" in text


def test_pdf_text_layer_fallback_when_no_form_fields() -> None:
    text = extract_text("loss_run.pdf", _build_plain_text_pdf())
    assert text is not None
    assert "Loss Run Summary" in text
    assert "Total Incurred 5yr: $20,700" in text


def test_pdf_form_fields_include_surrounding_static_title_text() -> None:
    """Regression: a fillable form's static title (e.g. drawn separately from the
    fields, as real ACORD PDFs do) must survive alongside the field values —
    dropping it silently broke content-based classification (no "acord" keyword
    left to match) on a real submitted PDF caught during live testing."""
    text = extract_text("acord_application.pdf", _build_form_pdf())
    assert text is not None
    assert "ACORD 125 - COMMERCIAL INSURANCE APPLICATION" in text
    assert "Named Insured: Delta Electric Services LLC" in text


def test_pdf_with_no_text_or_fields_returns_none() -> None:
    """A blank page (no form fields, no text layer) — honest 'nothing found',
    same as a genuinely scanned/image-only PDF would produce."""
    assert extract_text("scanned.pdf", _build_blank_pdf()) is None


def test_pdf_corrupt_bytes_never_raises_and_returns_none() -> None:
    assert extract_text("acord_application.pdf", b"\xff\xfe\xfd\xfc-not-a-real-pdf") is None


def test_docx_paragraphs_and_table_extracted() -> None:
    text = extract_text("financial_statement.docx", _build_docx())
    assert text is not None
    assert "Financial Statement Summary" in text
    assert "Annual Revenue: $5,800,000" in text
    assert "Total Assets: $2,400,000" in text


def test_docx_corrupt_bytes_returns_none() -> None:
    assert extract_text("financial_statement.docx", b"not a real docx") is None


def test_xlsx_two_column_sheet_and_wide_sheet_extracted() -> None:
    text = extract_text("sov_report.xlsx", _build_xlsx())
    assert text is not None
    assert "Total Insurable Value: 8500000" in text
    assert "Address | Building Value | Construction Type" in text
    assert "100 Main St | 1200000 | Frame" in text


def test_xlsx_corrupt_bytes_returns_none() -> None:
    assert extract_text("sov_report.xlsx", b"not a real xlsx") is None


def test_unsupported_extension_returns_none() -> None:
    assert extract_text("photo.png", b"\x89PNG\r\n\x1a\n") is None


@pytest.mark.parametrize("filename", ["acord.PDF", "loss_run.Pdf"])
def test_pdf_extension_matching_is_case_insensitive(filename: str) -> None:
    text = extract_text(filename, _build_plain_text_pdf())
    assert text is not None
    assert "Loss Run Summary" in text
